from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Iterable

from app.knowledge.models import KnowledgeDocument


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


def stable_document_id(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:24]


class DocumentLoader:
    def from_storage_docs(self, docs: Iterable[dict]) -> list[KnowledgeDocument]:
        documents: list[KnowledgeDocument] = []
        for item in docs:
            documents.append(
                KnowledgeDocument(
                    id=str(item["id"]),
                    title=str(item["title"]),
                    source_type=str(item.get("source_type", "document")),
                    domain=str(item.get("domain", "Business")),
                    tags=list(item.get("tags", [])),
                    content=str(item.get("content") or item.get("excerpt") or ""),
                    source_path=item.get("source_path"),
                    version=str(item.get("version", "v1")),
                    effective_date=item.get("effective_date"),
                    expires_at=item.get("expires_at"),
                    metadata=item.get("metadata", {}),
                )
            )
        return documents

    def from_folder(self, folder: Path, domain: str = "Enterprise Knowledge") -> list[KnowledgeDocument]:
        if not folder.exists():
            return []

        documents: list[KnowledgeDocument] = []
        for path in folder.rglob("*"):
            if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            content = self._read_file(path)
            if not content.strip():
                continue
            documents.append(
                KnowledgeDocument(
                    id=stable_document_id(str(path.resolve())),
                    title=path.stem.replace("_", " ").replace("-", " ").title(),
                    source_type=path.suffix.lower().lstrip("."),
                    domain=domain,
                    tags=[path.suffix.lower().lstrip("."), *[part.lower() for part in path.parts[-3:-1]]],
                    content=content,
                    source_path=str(path),
                    metadata={"file_name": path.name},
                )
            )
        return documents

    def _read_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md", ".markdown"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".pdf":
            return self._read_pdf(path)
        if suffix == ".docx":
            return self._read_docx(path)
        return ""

    def _read_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except Exception:
            return ""
        try:
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except Exception:
            return ""
        try:
            document = Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            return ""
